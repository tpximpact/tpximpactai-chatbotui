import asyncio
import copy
import json
import os
import re
import logging
import uuid
import docx
import io
from dotenv import load_dotenv

from quart import (
    Blueprint,
    Quart,
    websocket,
    jsonify,
    make_response,
    request,
    send_from_directory,
    render_template
)

from openai import AsyncAzureOpenAI
from azure.identity.aio import DefaultAzureCredential, get_bearer_token_provider
from backend.auth.auth_utils import get_authenticated_user_details
from backend.history.cosmosdbservice import CosmosConversationClient

from azure.monitor.opentelemetry import configure_azure_monitor
from azure.monitor.events.extension import track_event
from opentelemetry.trace import SpanKind


from opentelemetry import trace
from opentelemetry import metrics

import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter


from azure.core.credentials import AzureKeyCredential
from azure.identity.aio import ClientSecretCredential
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import *

from llama_index.core import StorageContext, VectorStoreIndex, Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.settings import Settings
from llama_index.llms.azure_openai import AzureOpenAI
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.vector_stores.azureaisearch import AzureAISearchVectorStore, IndexManagement
from openai import AzureOpenAI as BaseAzureOpenAI
# from llama_index.core import download_loader
from llama_index.readers.azstorage_blob import AzStorageBlobReader


configure_azure_monitor(connection_string= os.environ.get("APPLICATIONINSIGHTS_CONNECTION_STRING"))
tracer = trace.get_tracer(__name__)



from backend.utils import format_as_ndjson, format_stream_response, generateFilterString, generateSimpleFilterString, parse_multi_columns, format_non_streaming_response

bp = Blueprint("routes", __name__, static_folder="static", template_folder="static")

# Current minimum Azure OpenAI version supported
MINIMUM_SUPPORTED_AZURE_OPENAI_PREVIEW_API_VERSION="2024-02-15-preview"

# UI configuration (optional)
UI_TITLE = os.environ.get("UI_TITLE")
UI_LOGO = os.environ.get("UI_LOGO")
UI_CHAT_LOGO = os.environ.get("UI_CHAT_LOGO")
UI_CHAT_PRE_TITLE = os.environ.get("UI_CHAT_PRE_TITLE") or "Welcome to"
UI_CHAT_TITLE = os.environ.get("UI_CHAT_TITLE") or "ImpactAI"
UI_CHAT_DESCRIPTION = os.environ.get("UI_CHAT_DESCRIPTION")
UI_FAVICON = os.environ.get("UI_FAVICON") or "/favicon.ico"
UI_SHOW_SHARE_BUTTON = os.environ.get("UI_SHOW_SHARE_BUTTON", "true").lower() == "true"

def create_app():
    app = Quart(__name__)
    app.register_blueprint(bp)
    app.config["TEMPLATES_AUTO_RELOAD"] = True
    return app


@bp.route("/")
async def index():
    return await render_template("index.html", title=UI_TITLE, favicon=UI_FAVICON)

@bp.route("/favicon.ico")
async def favicon():
    return await bp.send_static_file("favicon.ico")

@bp.route("/assets/<path:path>")
async def assets(path):
    return await send_from_directory("static/assets", path)

load_dotenv()

# Debug settings
DEBUG = os.environ.get("DEBUG", "false")
if DEBUG.lower() == "true":
    logging.basicConfig(level=logging.DEBUG)


USER_AGENT = "GitHubSampleWebApp/AsyncAzureOpenAI/1.0.0"

# On Your Data Settings
DATASOURCE_TYPE = os.environ.get("DATASOURCE_TYPE", "AzureCognitiveSearch")
SEARCH_TOP_K = os.environ.get("SEARCH_TOP_K", 5)
SEARCH_STRICTNESS = os.environ.get("SEARCH_STRICTNESS", 3)
SEARCH_ENABLE_IN_DOMAIN = os.environ.get("SEARCH_ENABLE_IN_DOMAIN", "true")

# Document storage settings
AZURE_STORAGE_ACCOUNT = os.environ.get("AZURE_STORAGE_ACCOUNT")
AZURE_STORAGE_KEY = os.environ.get("AZURE_STORAGE_KEY")


# ACS Integration Settings
AZURE_SEARCH_SERVICE = os.environ.get("AZURE_SEARCH_SERVICE")
AZURE_SEARCH_INDEX = os.environ.get("AZURE_SEARCH_INDEX")
AZURE_SEARCH_KEY = os.environ.get("AZURE_SEARCH_KEY", None)
AZURE_SEARCH_USE_SEMANTIC_SEARCH = os.environ.get("AZURE_SEARCH_USE_SEMANTIC_SEARCH", "false")
AZURE_SEARCH_SEMANTIC_SEARCH_CONFIG = os.environ.get("AZURE_SEARCH_SEMANTIC_SEARCH_CONFIG", "default")
AZURE_SEARCH_TOP_K = os.environ.get("AZURE_SEARCH_TOP_K", SEARCH_TOP_K)
AZURE_SEARCH_ENABLE_IN_DOMAIN = os.environ.get("AZURE_SEARCH_ENABLE_IN_DOMAIN", SEARCH_ENABLE_IN_DOMAIN)
AZURE_SEARCH_CONTENT_COLUMNS = os.environ.get("AZURE_SEARCH_CONTENT_COLUMNS")
AZURE_SEARCH_FILENAME_COLUMN = os.environ.get("AZURE_SEARCH_FILENAME_COLUMN")
AZURE_SEARCH_TITLE_COLUMN = os.environ.get("AZURE_SEARCH_TITLE_COLUMN")
AZURE_SEARCH_URL_COLUMN = os.environ.get("AZURE_SEARCH_URL_COLUMN")
AZURE_SEARCH_VECTOR_COLUMNS = os.environ.get("AZURE_SEARCH_VECTOR_COLUMNS")
AZURE_SEARCH_QUERY_TYPE = os.environ.get("AZURE_SEARCH_QUERY_TYPE")
AZURE_SEARCH_PERMITTED_GROUPS_COLUMN = os.environ.get("AZURE_SEARCH_PERMITTED_GROUPS_COLUMN")
AZURE_SEARCH_STRICTNESS = os.environ.get("AZURE_SEARCH_STRICTNESS", SEARCH_STRICTNESS)

# AOAI Integration Settings
AZURE_OPENAI_RESOURCE = os.environ.get("AZURE_OPENAI_RESOURCE")
AZURE_OPENAI_MODEL = os.environ.get("AZURE_OPENAI_MODEL")
AZURE_OPENAI_ENDPOINT = os.environ.get("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.environ.get("AZURE_OPENAI_KEY")
AZURE_OPENAI_TEMPERATURE = os.environ.get("AZURE_OPENAI_TEMPERATURE", 0)
AZURE_OPENAI_TOP_P = os.environ.get("AZURE_OPENAI_TOP_P", 1.0)
AZURE_OPENAI_MAX_TOKENS = os.environ.get("AZURE_OPENAI_MAX_TOKENS", 1000)
AZURE_OPENAI_STOP_SEQUENCE = os.environ.get("AZURE_OPENAI_STOP_SEQUENCE")
AZURE_OPENAI_SYSTEM_MESSAGE = os.environ.get("AZURE_OPENAI_SYSTEM_MESSAGE", "You are an AI assistant that helps people find information.")
AZURE_OPENAI_PREVIEW_API_VERSION = os.environ.get("AZURE_OPENAI_PREVIEW_API_VERSION", MINIMUM_SUPPORTED_AZURE_OPENAI_PREVIEW_API_VERSION)
AZURE_OPENAI_STREAM = os.environ.get("AZURE_OPENAI_STREAM", "true")
AZURE_OPENAI_MODEL_NAME = os.environ.get("AZURE_OPENAI_MODEL_NAME", "gpt-35-turbo-16k") # Name of the model, e.g. 'gpt-35-turbo-16k' or 'gpt-4'
AZURE_OPENAI_EMBEDDING_ENDPOINT = os.environ.get("AZURE_OPENAI_EMBEDDING_ENDPOINT")
AZURE_OPENAI_EMBEDDING_KEY = os.environ.get("AZURE_OPENAI_EMBEDDING_KEY")
AZURE_OPENAI_EMBEDDING_NAME = os.environ.get("AZURE_OPENAI_EMBEDDING_NAME", "")
AZURE_OPENAI_EMBEDDING_MODEL_NAME = os.environ.get("AZURE_OPENAI_EMBEDDING_MODEL_NAME")

# CosmosDB Mongo vcore vector db Settings
AZURE_COSMOSDB_MONGO_VCORE_CONNECTION_STRING = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_CONNECTION_STRING")  #This has to be secure string
AZURE_COSMOSDB_MONGO_VCORE_DATABASE = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_DATABASE")
AZURE_COSMOSDB_MONGO_VCORE_CONTAINER = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_CONTAINER")
AZURE_COSMOSDB_MONGO_VCORE_INDEX = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_INDEX")
AZURE_COSMOSDB_MONGO_VCORE_TOP_K = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_TOP_K", AZURE_SEARCH_TOP_K)
AZURE_COSMOSDB_MONGO_VCORE_STRICTNESS = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_STRICTNESS", AZURE_SEARCH_STRICTNESS)  
AZURE_COSMOSDB_MONGO_VCORE_ENABLE_IN_DOMAIN = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_ENABLE_IN_DOMAIN", AZURE_SEARCH_ENABLE_IN_DOMAIN)
AZURE_COSMOSDB_MONGO_VCORE_CONTENT_COLUMNS = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_CONTENT_COLUMNS", "")
AZURE_COSMOSDB_MONGO_VCORE_FILENAME_COLUMN = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_FILENAME_COLUMN")
AZURE_COSMOSDB_MONGO_VCORE_TITLE_COLUMN = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_TITLE_COLUMN")
AZURE_COSMOSDB_MONGO_VCORE_URL_COLUMN = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_URL_COLUMN")
AZURE_COSMOSDB_MONGO_VCORE_VECTOR_COLUMNS = os.environ.get("AZURE_COSMOSDB_MONGO_VCORE_VECTOR_COLUMNS")

SHOULD_STREAM = True if AZURE_OPENAI_STREAM.lower() == "true" else False

# Chat History CosmosDB Integration Settings
AZURE_COSMOSDB_DATABASE = os.environ.get("AZURE_COSMOSDB_DATABASE")
AZURE_COSMOSDB_ACCOUNT = os.environ.get("AZURE_COSMOSDB_ACCOUNT")
AZURE_COSMOSDB_CONVERSATIONS_CONTAINER = os.environ.get("AZURE_COSMOSDB_CONVERSATIONS_CONTAINER")
AZURE_COSMOSDB_ACCOUNT_KEY = os.environ.get("AZURE_COSMOSDB_ACCOUNT_KEY")
AZURE_COSMOSDB_ENABLE_FEEDBACK = os.environ.get("AZURE_COSMOSDB_ENABLE_FEEDBACK", "false").lower() == "true"

# Elasticsearch Integration Settings
ELASTICSEARCH_ENDPOINT = os.environ.get("ELASTICSEARCH_ENDPOINT")
ELASTICSEARCH_ENCODED_API_KEY = os.environ.get("ELASTICSEARCH_ENCODED_API_KEY")
ELASTICSEARCH_INDEX = os.environ.get("ELASTICSEARCH_INDEX")
ELASTICSEARCH_QUERY_TYPE = os.environ.get("ELASTICSEARCH_QUERY_TYPE", "simple")
ELASTICSEARCH_TOP_K = os.environ.get("ELASTICSEARCH_TOP_K", SEARCH_TOP_K)
ELASTICSEARCH_ENABLE_IN_DOMAIN = os.environ.get("ELASTICSEARCH_ENABLE_IN_DOMAIN", SEARCH_ENABLE_IN_DOMAIN)
ELASTICSEARCH_CONTENT_COLUMNS = os.environ.get("ELASTICSEARCH_CONTENT_COLUMNS")
ELASTICSEARCH_FILENAME_COLUMN = os.environ.get("ELASTICSEARCH_FILENAME_COLUMN")
ELASTICSEARCH_TITLE_COLUMN = os.environ.get("ELASTICSEARCH_TITLE_COLUMN")
ELASTICSEARCH_URL_COLUMN = os.environ.get("ELASTICSEARCH_URL_COLUMN")
ELASTICSEARCH_VECTOR_COLUMNS = os.environ.get("ELASTICSEARCH_VECTOR_COLUMNS")
ELASTICSEARCH_STRICTNESS = os.environ.get("ELASTICSEARCH_STRICTNESS", SEARCH_STRICTNESS)
ELASTICSEARCH_EMBEDDING_MODEL_ID = os.environ.get("ELASTICSEARCH_EMBEDDING_MODEL_ID")

# Pinecone Integration Settings
PINECONE_ENVIRONMENT = os.environ.get("PINECONE_ENVIRONMENT")
PINECONE_API_KEY = os.environ.get("PINECONE_API_KEY")
PINECONE_INDEX_NAME = os.environ.get("PINECONE_INDEX_NAME")
PINECONE_TOP_K = os.environ.get("PINECONE_TOP_K", SEARCH_TOP_K)
PINECONE_STRICTNESS = os.environ.get("PINECONE_STRICTNESS", SEARCH_STRICTNESS)  
PINECONE_ENABLE_IN_DOMAIN = os.environ.get("PINECONE_ENABLE_IN_DOMAIN", SEARCH_ENABLE_IN_DOMAIN)
PINECONE_CONTENT_COLUMNS = os.environ.get("PINECONE_CONTENT_COLUMNS", "")
PINECONE_FILENAME_COLUMN = os.environ.get("PINECONE_FILENAME_COLUMN")
PINECONE_TITLE_COLUMN = os.environ.get("PINECONE_TITLE_COLUMN")
PINECONE_URL_COLUMN = os.environ.get("PINECONE_URL_COLUMN")
PINECONE_VECTOR_COLUMNS = os.environ.get("PINECONE_VECTOR_COLUMNS")

# Azure AI MLIndex Integration Settings - for use with MLIndex data assets created in Azure AI Studio
AZURE_MLINDEX_NAME = os.environ.get("AZURE_MLINDEX_NAME")
AZURE_MLINDEX_VERSION = os.environ.get("AZURE_MLINDEX_VERSION")
AZURE_ML_PROJECT_RESOURCE_ID = os.environ.get("AZURE_ML_PROJECT_RESOURCE_ID") # /subscriptions/{sub ID}/resourceGroups/{rg name}/providers/Microsoft.MachineLearningServices/workspaces/{AML project name}
AZURE_MLINDEX_TOP_K = os.environ.get("AZURE_MLINDEX_TOP_K", SEARCH_TOP_K)
AZURE_MLINDEX_STRICTNESS = os.environ.get("AZURE_MLINDEX_STRICTNESS", SEARCH_STRICTNESS)  
AZURE_MLINDEX_ENABLE_IN_DOMAIN = os.environ.get("AZURE_MLINDEX_ENABLE_IN_DOMAIN", SEARCH_ENABLE_IN_DOMAIN)
AZURE_MLINDEX_CONTENT_COLUMNS = os.environ.get("AZURE_MLINDEX_CONTENT_COLUMNS", "")
AZURE_MLINDEX_FILENAME_COLUMN = os.environ.get("AZURE_MLINDEX_FILENAME_COLUMN")
AZURE_MLINDEX_TITLE_COLUMN = os.environ.get("AZURE_MLINDEX_TITLE_COLUMN")
AZURE_MLINDEX_URL_COLUMN = os.environ.get("AZURE_MLINDEX_URL_COLUMN")
AZURE_MLINDEX_VECTOR_COLUMNS = os.environ.get("AZURE_MLINDEX_VECTOR_COLUMNS")
AZURE_MLINDEX_QUERY_TYPE = os.environ.get("AZURE_MLINDEX_QUERY_TYPE")


# Frontend Settings via Environment Variables
AUTH_ENABLED = os.environ.get("AUTH_ENABLED", "true").lower() == "true"
CHAT_HISTORY_ENABLED = AZURE_COSMOSDB_ACCOUNT and AZURE_COSMOSDB_DATABASE and AZURE_COSMOSDB_CONVERSATIONS_CONTAINER
SANITIZE_ANSWER = os.environ.get("SANITIZE_ANSWER", "false").lower() == "true"
frontend_settings = { 
    "auth_enabled": AUTH_ENABLED, 
    "feedback_enabled": AZURE_COSMOSDB_ENABLE_FEEDBACK and CHAT_HISTORY_ENABLED,
    "ui": {
        "title": UI_TITLE,
        "logo": UI_LOGO,
        "chat_logo": UI_CHAT_LOGO or UI_LOGO,
        "chat_pre_title": UI_CHAT_PRE_TITLE,
        "chat_title": UI_CHAT_TITLE,
        "chat_description": UI_CHAT_DESCRIPTION,
        "show_share_button": UI_SHOW_SHARE_BUTTON
    },
    "sanitize_answer": SANITIZE_ANSWER
}

def should_use_data():
    global DATASOURCE_TYPE
    if AZURE_SEARCH_SERVICE and AZURE_SEARCH_INDEX:
        DATASOURCE_TYPE = "AzureCognitiveSearch"
        logging.debug("Using Azure Cognitive Search")
        return True
    
    if AZURE_COSMOSDB_MONGO_VCORE_DATABASE and AZURE_COSMOSDB_MONGO_VCORE_CONTAINER and AZURE_COSMOSDB_MONGO_VCORE_INDEX and AZURE_COSMOSDB_MONGO_VCORE_CONNECTION_STRING:
        DATASOURCE_TYPE = "AzureCosmosDB"
        logging.debug("Using Azure CosmosDB Mongo vcore")
        return True
    
    if ELASTICSEARCH_ENDPOINT and ELASTICSEARCH_ENCODED_API_KEY and ELASTICSEARCH_INDEX:
        DATASOURCE_TYPE = "Elasticsearch"
        logging.debug("Using Elasticsearch")
        return True
    
    if PINECONE_ENVIRONMENT and PINECONE_API_KEY and PINECONE_INDEX_NAME:
        DATASOURCE_TYPE = "Pinecone"
        logging.debug("Using Pinecone")
        return True
    
    if AZURE_MLINDEX_NAME and AZURE_MLINDEX_VERSION and AZURE_ML_PROJECT_RESOURCE_ID:
        DATASOURCE_TYPE = "AzureMLIndex"
        logging.debug("Using Azure ML Index")
        return True

    return False

SHOULD_USE_DATA = should_use_data()

# Initialize Azure OpenAI Client
def init_openai_client(use_data=SHOULD_USE_DATA):
    azure_openai_client = None
    try:
        # API version check
        if AZURE_OPENAI_PREVIEW_API_VERSION < MINIMUM_SUPPORTED_AZURE_OPENAI_PREVIEW_API_VERSION:
            raise Exception(f"The minimum supported Azure OpenAI preview API version is '{MINIMUM_SUPPORTED_AZURE_OPENAI_PREVIEW_API_VERSION}'")
        
        # Endpoint
        if not AZURE_OPENAI_ENDPOINT and not AZURE_OPENAI_RESOURCE:
            raise Exception("AZURE_OPENAI_ENDPOINT or AZURE_OPENAI_RESOURCE is required")
        
        endpoint = AZURE_OPENAI_ENDPOINT if AZURE_OPENAI_ENDPOINT else f"https://{AZURE_OPENAI_RESOURCE}.openai.azure.com/"
        
        # Authentication
        aoai_api_key = AZURE_OPENAI_KEY
        ad_token_provider = None
        if not aoai_api_key:
            logging.debug("No AZURE_OPENAI_KEY found, using Azure AD auth")
            ad_token_provider = get_bearer_token_provider(DefaultAzureCredential(), "https://cognitiveservices.azure.com/.default")

        # Deployment
        deployment = AZURE_OPENAI_MODEL
        if not deployment:
            raise Exception("AZURE_OPENAI_MODEL is required")

        # Default Headers
        default_headers = {
            'x-ms-useragent': USER_AGENT
        }

        azure_openai_client = AsyncAzureOpenAI(
            api_version=AZURE_OPENAI_PREVIEW_API_VERSION,
            api_key=aoai_api_key,
            azure_ad_token_provider=ad_token_provider,
            default_headers=default_headers,
            azure_endpoint=endpoint
        )
            
        return azure_openai_client
    except Exception as e:
        logging.exception("Exception in Azure OpenAI initialization", e)
        azure_openai_client = None
        raise e


def init_cosmosdb_client():
    logging.debug("Initializing CosmosDB client")
    cosmos_conversation_client = None
    if CHAT_HISTORY_ENABLED:
        logging.debug("COSMOSDB ENABLED")
        try:
            cosmos_endpoint = f'https://{AZURE_COSMOSDB_ACCOUNT}.documents.azure.com:443/'

            if not AZURE_COSMOSDB_ACCOUNT_KEY:
                credential = DefaultAzureCredential()
            else:
                credential = AZURE_COSMOSDB_ACCOUNT_KEY

            cosmos_conversation_client = CosmosConversationClient(
                cosmosdb_endpoint=cosmos_endpoint, 
                credential=credential, 
                database_name=AZURE_COSMOSDB_DATABASE,
                container_name=AZURE_COSMOSDB_CONVERSATIONS_CONTAINER,
                enable_message_feedback=AZURE_COSMOSDB_ENABLE_FEEDBACK
            )
        except Exception as e:
            logging.exception("Exception in CosmosDB initialization", e)
            cosmos_conversation_client = None
            raise e
    else:
        logging.debug("CosmosDB not configured")
        
    return cosmos_conversation_client

def init_container_client(storage_container_name):

    storage_account_url = f"https://{AZURE_STORAGE_ACCOUNT}.blob.core.windows.net/"

    blob_service_client = BlobServiceClient(
    account_url=storage_account_url, 
    credential=AZURE_STORAGE_KEY
    )

    container_client = blob_service_client.get_container_client(storage_container_name)
    return container_client

def init_search_client():
    search_client = None
    try:
        search_client = SearchClient(
            endpoint=f"https://{AZURE_SEARCH_SERVICE}.search.windows.net/",
            index_name=AZURE_SEARCH_INDEX,
            credential=AzureKeyCredential(AZURE_SEARCH_KEY)
        )
    except Exception as e:
        logging.exception("Exception in Azure Cognitive Search initialization", e)
        search_client = None
        raise e
        
    return search_client

def get_configured_data_source(user_id, filenames):
    data_source = {}
    query_type = "simple"
    if DATASOURCE_TYPE == "AzureCognitiveSearch":
        # Set query type
        if AZURE_SEARCH_QUERY_TYPE:
            query_type = AZURE_SEARCH_QUERY_TYPE
        elif AZURE_SEARCH_USE_SEMANTIC_SEARCH.lower() == "true" and AZURE_SEARCH_SEMANTIC_SEARCH_CONFIG:
            query_type = "semantic"

        # Set filter
        filter = generateSimpleFilterString(user_id, filenames)
        userToken = None
        if AZURE_SEARCH_PERMITTED_GROUPS_COLUMN:
            userToken = request.headers.get('X-MS-TOKEN-AAD-ACCESS-TOKEN', "")
            logging.debug(f"USER TOKEN is {'present' if userToken else 'not present'}")
            if not userToken:
                raise Exception("Document-level access control is enabled, but user access token could not be fetched.")

            filter = generateFilterString(userToken)
            logging.debug(f"FILTER: {filter}")
        # Set authentication
        authentication = {}
        if AZURE_SEARCH_KEY:
            authentication = {
                "type": "api_key",
                "api_key": AZURE_SEARCH_KEY
            }
        else:
            # If key is not provided, assume AOAI resource identity has been granted access to the search service
            logging.debug("Using system-assigned managed identity for Azure Cognitive Search")
            authentication = {
                "type": "system_assigned_managed_identity"
            }

        data_source = {
                "type": "azure_search",
                "parameters": {
                    "endpoint": f"https://{AZURE_SEARCH_SERVICE}.search.windows.net",
                    "authentication": authentication,
                    "index_name": AZURE_SEARCH_INDEX,
                    "fields_mapping": {
                        "content_fields": parse_multi_columns(AZURE_SEARCH_CONTENT_COLUMNS) if AZURE_SEARCH_CONTENT_COLUMNS else [],
                        "title_field": AZURE_SEARCH_TITLE_COLUMN if AZURE_SEARCH_TITLE_COLUMN else None,
                        "url_field": AZURE_SEARCH_URL_COLUMN if AZURE_SEARCH_URL_COLUMN else None,
                        "filepath_field": AZURE_SEARCH_FILENAME_COLUMN if AZURE_SEARCH_FILENAME_COLUMN else None,
                        "vector_fields": parse_multi_columns(AZURE_SEARCH_VECTOR_COLUMNS) if AZURE_SEARCH_VECTOR_COLUMNS else []
                    },
                    "in_scope": True if AZURE_SEARCH_ENABLE_IN_DOMAIN.lower() == "true" else False,
                    "top_n_documents": int(AZURE_SEARCH_TOP_K) if AZURE_SEARCH_TOP_K else int(SEARCH_TOP_K),
                    "query_type": query_type,
                    "semantic_configuration": AZURE_SEARCH_SEMANTIC_SEARCH_CONFIG if AZURE_SEARCH_SEMANTIC_SEARCH_CONFIG else "",
                    "role_information": AZURE_OPENAI_SYSTEM_MESSAGE,
                    "filter": filter,
                    "strictness": int(AZURE_SEARCH_STRICTNESS) if AZURE_SEARCH_STRICTNESS else int(SEARCH_STRICTNESS)
                }
            }
    elif DATASOURCE_TYPE == "AzureCosmosDB":
        query_type = "vector"

        data_source = {
                "type": "azure_cosmos_db",
                "parameters": {
                    "authentication": {
                        "type": "connection_string",
                        "connection_string": AZURE_COSMOSDB_MONGO_VCORE_CONNECTION_STRING
                    },
                    "index_name": AZURE_COSMOSDB_MONGO_VCORE_INDEX,
                    "database_name": AZURE_COSMOSDB_MONGO_VCORE_DATABASE,
                    "container_name": AZURE_COSMOSDB_MONGO_VCORE_CONTAINER,                    
                    "fields_mapping": {
                        "content_fields": parse_multi_columns(AZURE_COSMOSDB_MONGO_VCORE_CONTENT_COLUMNS) if AZURE_COSMOSDB_MONGO_VCORE_CONTENT_COLUMNS else [],
                        "title_field": AZURE_COSMOSDB_MONGO_VCORE_TITLE_COLUMN if AZURE_COSMOSDB_MONGO_VCORE_TITLE_COLUMN else None,
                        "url_field": AZURE_COSMOSDB_MONGO_VCORE_URL_COLUMN if AZURE_COSMOSDB_MONGO_VCORE_URL_COLUMN else None,
                        "filepath_field": AZURE_COSMOSDB_MONGO_VCORE_FILENAME_COLUMN if AZURE_COSMOSDB_MONGO_VCORE_FILENAME_COLUMN else None,
                        "vector_fields": parse_multi_columns(AZURE_COSMOSDB_MONGO_VCORE_VECTOR_COLUMNS) if AZURE_COSMOSDB_MONGO_VCORE_VECTOR_COLUMNS else []
                    },
                    "in_scope": True if AZURE_COSMOSDB_MONGO_VCORE_ENABLE_IN_DOMAIN.lower() == "true" else False,
                    "top_n_documents": int(AZURE_COSMOSDB_MONGO_VCORE_TOP_K) if AZURE_COSMOSDB_MONGO_VCORE_TOP_K else int(SEARCH_TOP_K),
                    "strictness": int(AZURE_COSMOSDB_MONGO_VCORE_STRICTNESS) if AZURE_COSMOSDB_MONGO_VCORE_STRICTNESS else int(SEARCH_STRICTNESS),
                    "query_type": query_type,
                    "role_information": AZURE_OPENAI_SYSTEM_MESSAGE
                }
            }
    elif DATASOURCE_TYPE == "Elasticsearch":
        if ELASTICSEARCH_QUERY_TYPE:
            query_type = ELASTICSEARCH_QUERY_TYPE

        data_source = {
            "type": "elasticsearch",
            "parameters": {
                "endpoint": ELASTICSEARCH_ENDPOINT,
                "authentication": {
                    "type": "encoded_api_key",
                    "encoded_api_key": ELASTICSEARCH_ENCODED_API_KEY
                },
                "index_name": ELASTICSEARCH_INDEX,
                "fields_mapping": {
                    "content_fields": parse_multi_columns(ELASTICSEARCH_CONTENT_COLUMNS) if ELASTICSEARCH_CONTENT_COLUMNS else [],
                    "title_field": ELASTICSEARCH_TITLE_COLUMN if ELASTICSEARCH_TITLE_COLUMN else None,
                    "url_field": ELASTICSEARCH_URL_COLUMN if ELASTICSEARCH_URL_COLUMN else None,
                    "filepath_field": ELASTICSEARCH_FILENAME_COLUMN if ELASTICSEARCH_FILENAME_COLUMN else None,
                    "vector_fields": parse_multi_columns(ELASTICSEARCH_VECTOR_COLUMNS) if ELASTICSEARCH_VECTOR_COLUMNS else []
                },
                "in_scope": True if ELASTICSEARCH_ENABLE_IN_DOMAIN.lower() == "true" else False,
                "top_n_documents": int(ELASTICSEARCH_TOP_K) if ELASTICSEARCH_TOP_K else int(SEARCH_TOP_K),
                "query_type": query_type,
                "role_information": AZURE_OPENAI_SYSTEM_MESSAGE,
                "strictness": int(ELASTICSEARCH_STRICTNESS) if ELASTICSEARCH_STRICTNESS else int(SEARCH_STRICTNESS)
            }
        }
    elif DATASOURCE_TYPE == "AzureMLIndex":
        if AZURE_MLINDEX_QUERY_TYPE:
            query_type = AZURE_MLINDEX_QUERY_TYPE

        data_source = {
            "type": "azure_ml_index",
            "parameters": {
                "name": AZURE_MLINDEX_NAME,
                "version": AZURE_MLINDEX_VERSION,
                "project_resource_id": AZURE_ML_PROJECT_RESOURCE_ID,
                "fieldsMapping": {
                    "content_fields": parse_multi_columns(AZURE_MLINDEX_CONTENT_COLUMNS) if AZURE_MLINDEX_CONTENT_COLUMNS else [],
                    "title_field": AZURE_MLINDEX_TITLE_COLUMN if AZURE_MLINDEX_TITLE_COLUMN else None,
                    "url_field": AZURE_MLINDEX_URL_COLUMN if AZURE_MLINDEX_URL_COLUMN else None,
                    "filepath_field": AZURE_MLINDEX_FILENAME_COLUMN if AZURE_MLINDEX_FILENAME_COLUMN else None,
                    "vector_fields": parse_multi_columns(AZURE_MLINDEX_VECTOR_COLUMNS) if AZURE_MLINDEX_VECTOR_COLUMNS else []
                },
                "in_scope": True if AZURE_MLINDEX_ENABLE_IN_DOMAIN.lower() == "true" else False,
                "top_n_documents": int(AZURE_MLINDEX_TOP_K) if AZURE_MLINDEX_TOP_K else int(SEARCH_TOP_K),
                "query_type": query_type,
                "role_information": AZURE_OPENAI_SYSTEM_MESSAGE,
                "strictness": int(AZURE_MLINDEX_STRICTNESS) if AZURE_MLINDEX_STRICTNESS else int(SEARCH_STRICTNESS)
            }
        }
    elif DATASOURCE_TYPE == "Pinecone":
        query_type = "vector"

        data_source = {
            "type": "pinecone",
            "parameters": {
                "environment": PINECONE_ENVIRONMENT,
                "authentication": {
                    "type": "api_key",
                    "key": PINECONE_API_KEY
                },
                "index_name": PINECONE_INDEX_NAME,
                "fields_mapping": {
                    "content_fields": parse_multi_columns(PINECONE_CONTENT_COLUMNS) if PINECONE_CONTENT_COLUMNS else [],
                    "title_field": PINECONE_TITLE_COLUMN if PINECONE_TITLE_COLUMN else None,
                    "url_field": PINECONE_URL_COLUMN if PINECONE_URL_COLUMN else None,
                    "filepath_field": PINECONE_FILENAME_COLUMN if PINECONE_FILENAME_COLUMN else None,
                    "vector_fields": parse_multi_columns(PINECONE_VECTOR_COLUMNS) if PINECONE_VECTOR_COLUMNS else []
                },
                "in_scope": True if PINECONE_ENABLE_IN_DOMAIN.lower() == "true" else False,
                "top_n_documents": int(PINECONE_TOP_K) if PINECONE_TOP_K else int(SEARCH_TOP_K),
                "strictness": int(PINECONE_STRICTNESS) if PINECONE_STRICTNESS else int(SEARCH_STRICTNESS),
                "query_type": query_type,
                "role_information": AZURE_OPENAI_SYSTEM_MESSAGE,
            }
        }
    else:
        raise Exception(f"DATASOURCE_TYPE is not configured or unknown: {DATASOURCE_TYPE}")

    if "vector" in query_type.lower() and DATASOURCE_TYPE != "AzureMLIndex":
        embeddingDependency = {}
        if AZURE_OPENAI_EMBEDDING_NAME:
            embeddingDependency = {
                "type": "deployment_name",
                "deployment_name": AZURE_OPENAI_EMBEDDING_NAME
            }
        elif AZURE_OPENAI_EMBEDDING_ENDPOINT and AZURE_OPENAI_EMBEDDING_KEY:
            embeddingDependency = {
                "type": "endpoint",
                "endpoint": AZURE_OPENAI_EMBEDDING_ENDPOINT,
                "authentication": {
                    "type": "api_key",
                    "key": AZURE_OPENAI_EMBEDDING_KEY
                }
            }
        elif DATASOURCE_TYPE == "Elasticsearch" and ELASTICSEARCH_EMBEDDING_MODEL_ID:
            embeddingDependency = {
                "type": "model_id",
                "model_id": ELASTICSEARCH_EMBEDDING_MODEL_ID
            }
        else:
            raise Exception(f"Vector query type ({query_type}) is selected for data source type {DATASOURCE_TYPE} but no embedding dependency is configured")
        data_source["parameters"]["embedding_dependency"] = embeddingDependency

    return data_source

def prepare_model_args(request_body):
    logging.debug("PREPARING MODEL ARGS")
    try:
        authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    except Exception as e:
        logging.exception("Exception in get_authenticated_user_details")
        raise e
    user_id = authenticated_user['user_principal_id']
    request_messages = request_body.get("messages", [])
    request_filenames = request_body.get("filenames", [])
    messages = []
    if request_filenames == []:
        messages = [
            {
                "role": "system",
                "content": AZURE_OPENAI_SYSTEM_MESSAGE
            }
        ]

    for message in request_messages:
        if message:
            messages.append({
                "role": message["role"] ,
                "content": message["content"]
            })

    model_args = {
        "messages": messages,
        "temperature": float(AZURE_OPENAI_TEMPERATURE),
        "max_tokens": int(AZURE_OPENAI_MAX_TOKENS),
        "top_p": float(AZURE_OPENAI_TOP_P),
        "stop": parse_multi_columns(AZURE_OPENAI_STOP_SEQUENCE) if AZURE_OPENAI_STOP_SEQUENCE else None,
        "stream": SHOULD_STREAM,
        "model": AZURE_OPENAI_MODEL,
    }

    if len(request_filenames) > 0:
        model_args["extra_body"] = {
            "data_sources": [get_configured_data_source(user_id, request_filenames)]
        }

    model_args_clean = copy.deepcopy(model_args)
    if model_args_clean.get("extra_body"):
        secret_params = ["key", "connection_string", "embedding_key", "encoded_api_key", "api_key"]
        for secret_param in secret_params:
            if model_args_clean["extra_body"]["data_sources"][0]["parameters"].get(secret_param):
                model_args_clean["extra_body"]["data_sources"][0]["parameters"][secret_param] = "*****"
        authentication = model_args_clean["extra_body"]["data_sources"][0]["parameters"].get("authentication", {})
        for field in authentication:
            if field in secret_params:
                model_args_clean["extra_body"]["data_sources"][0]["parameters"]["authentication"][field] = "*****"
        embeddingDependency = model_args_clean["extra_body"]["data_sources"][0]["parameters"].get("embedding_dependency", {})
        if "authentication" in embeddingDependency:
            for field in embeddingDependency["authentication"]:
                if field in secret_params:
                    model_args_clean["extra_body"]["data_sources"][0]["parameters"]["embedding_dependency"]["authentication"][field] = "*****"
        
    logging.debug(f"REQUEST BODY: {json.dumps(model_args_clean, indent=4)}")
    
    return model_args

async def send_chat_request(request):
    model_args = prepare_model_args(request)

    try:
        azure_openai_client = init_openai_client()
        response = await azure_openai_client.chat.completions.create(**model_args)


    except Exception as e:
        logging.exception("Exception in send_chat_request")
        raise e

    return response

async def complete_chat_request(request_body):
    response = await send_chat_request(request_body)
    history_metadata = request_body.get("history_metadata", {})

    return format_non_streaming_response(response, history_metadata)

async def stream_chat_request(request_body):
    response = await send_chat_request(request_body)
    history_metadata = request_body.get("history_metadata", {})

    async def generate():
        async for completionChunk in response:
            yield format_stream_response(completionChunk, history_metadata)

    return generate()

async def conversation_internal(request_body):
    try:
        if SHOULD_STREAM:
            result = await stream_chat_request(request_body)
            response = await make_response(format_as_ndjson(result))
            response.timeout = None
            response.mimetype = "application/json-lines"
            return response
        else:
            result = await complete_chat_request(request_body)
            return jsonify(result)
    
    except Exception as ex:
        logging.exception(ex)
        if hasattr(ex, "status_code"):
            return jsonify({"error": str(ex)}), ex.status_code
        else:
            return jsonify({"error": str(ex)}), 500


meter = metrics.get_meter_provider().get_meter("otel_azure_monitor_counter_demo")
counter = meter.create_counter("counter")

@bp.route("/conversation", methods=["POST"])
async def conversation():

    if not request.is_json:
        return jsonify({"error": "request must be json"}), 415
    request_json = await request.get_json()

    return await conversation_internal(request_json)

@bp.route("/frontend_settings", methods=["GET"])  
def get_frontend_settings():
    try:
        return jsonify(frontend_settings), 200
    except Exception as e:
        logging.exception("Exception in /frontend_settings")
        return jsonify({"error": str(e)}), 500  

## Conversation History API ##
@bp.route("/history/generate", methods=["POST"])
async def add_conversation():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)

    try:
        # make sure cosmos is configured
        cosmos_conversation_client = init_cosmosdb_client()
        if not cosmos_conversation_client:
            raise Exception("CosmosDB is not configured or not working")

        # check for the conversation_id, if the conversation is not set, we will create a new one
        history_metadata = {}
        if not conversation_id:
            title = await generate_title(request_json["messages"])
            conversation_dict = await cosmos_conversation_client.create_conversation(user_id=user_id, title=title)
            conversation_id = conversation_dict['id']
            history_metadata['title'] = title
            history_metadata['date'] = conversation_dict['createdAt']
            
        ## Format the incoming message object in the "chat/completions" messages format
        ## then write it to the conversation history in cosmos
        messages = request_json["messages"]
        if len(messages) > 0 and messages[-1]['role'] == "user":
            createdMessageValue = await cosmos_conversation_client.create_message(
                uuid=str(uuid.uuid4()),
                conversation_id=conversation_id,
                user_id=user_id,
                input_message=messages[-1],
                hidden=messages[-1]['hidden']
            )
            if createdMessageValue == "Conversation not found":
                raise Exception("Conversation not found for the given conversation ID: " + conversation_id + ".")
        else:
            raise Exception("No user message found")
        
        await cosmos_conversation_client.cosmosdb_client.close()
        
        # Submit request to Chat Completions for response
        request_body = await request.get_json()
        history_metadata['conversation_id'] = conversation_id
        request_body['history_metadata'] = history_metadata
        return await conversation_internal(request_body)
       
    except Exception as e:
        logging.exception("Exception in /history/generate")
        return jsonify({"error": str(e)}), 500


@bp.route("/history/update", methods=["POST"])
async def update_conversation():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)

    try:
        # make sure cosmos is configured
        cosmos_conversation_client = init_cosmosdb_client()
        if not cosmos_conversation_client:
            raise Exception("CosmosDB is not configured or not working")

        # check for the conversation_id, if the conversation is not set, we will create a new one
        if not conversation_id:
            raise Exception("No conversation_id found")
            
        ## Format the incoming message object in the "chat/completions" messages format
        ## then write it to the conversation history in cosmos
        messages = request_json["messages"]
        if len(messages) > 0 and messages[-1]['role'] == "assistant":
            if len(messages) > 1 and messages[-2].get('role', None) == "tool":
                # write the tool message first
                await cosmos_conversation_client.create_message(
                    uuid=str(uuid.uuid4()),
                    conversation_id=conversation_id,
                    user_id=user_id,
                    input_message=messages[-2]
                )
            # write the assistant message
            await cosmos_conversation_client.create_message(
                uuid=messages[-1]['id'],
                conversation_id=conversation_id,
                user_id=user_id,
                input_message=messages[-1]
            )
        else:
            raise Exception("No bot messages found")
        
        # Submit request to Chat Completions for response
        await cosmos_conversation_client.cosmosdb_client.close()
        response = {'success': True}
        return jsonify(response), 200
       
    except Exception as e:
        logging.exception("Exception in /history/update")
        return jsonify({"error": str(e)}), 500

@bp.route("/history/message_feedback", methods=["POST"])
async def update_message():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']
    cosmos_conversation_client = init_cosmosdb_client()

    ## check request for message_id
    request_json = await request.get_json()
    message_id = request_json.get('message_id', None)
    message_feedback = request_json.get("message_feedback", None)
    try:
        if not message_id:
            return jsonify({"error": "message_id is required"}), 400
        
        if not message_feedback:
            return jsonify({"error": "message_feedback is required"}), 400
        
        ## update the message in cosmos
        updated_message = await cosmos_conversation_client.update_message_feedback(user_id, message_id, message_feedback)
        if updated_message:
            return jsonify({"message": f"Successfully updated message with feedback {message_feedback}", "message_id": message_id}), 200
        else:
            return jsonify({"error": f"Unable to update message {message_id}. It either does not exist or the user does not have access to it."}), 404
        
    except Exception as e:
        logging.exception("Exception in /history/message_feedback")
        return jsonify({"error": str(e)}), 500


@bp.route("/history/delete", methods=["DELETE"])
async def delete_conversation():

    ## get the user id from the request headers
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']
    
    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)

    try: 
        if not conversation_id:
            return jsonify({"error": "conversation_id is required"}), 400
        
        ## make sure cosmos is configured
        cosmos_conversation_client = init_cosmosdb_client()
        if not cosmos_conversation_client:
            raise Exception("CosmosDB is not configured or not working")

        ## delete the conversation messages from cosmos first
        deleted_messages = await cosmos_conversation_client.delete_messages(conversation_id, user_id)

        ## Now delete the conversation 
        deleted_conversation = await cosmos_conversation_client.delete_conversation(user_id, conversation_id)

        await cosmos_conversation_client.cosmosdb_client.close()

        return jsonify({"message": "Successfully deleted conversation and messages", "conversation_id": conversation_id}), 200
    except Exception as e:
        logging.exception("Exception in /history/delete")
        return jsonify({"error": str(e)}), 500


@bp.route("/history/list", methods=["GET"])
async def list_conversations():
    offset = request.args.get("offset", 0)
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    ## make sure cosmos is configured
    cosmos_conversation_client = init_cosmosdb_client()
    if not cosmos_conversation_client:
        raise Exception("CosmosDB is not configured or not working")

    ## get the conversations from cosmos
    conversations = await cosmos_conversation_client.get_conversations(user_id, offset=offset, limit=25)
    await cosmos_conversation_client.cosmosdb_client.close()
    if not isinstance(conversations, list):
        return jsonify({"error": f"No conversations for {user_id} were found"}), 404

    ## return the conversation ids

    return jsonify(conversations), 200


@bp.route("/history/read", methods=["POST"])
async def get_conversation():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)
    
    if not conversation_id:
        return jsonify({"error": "conversation_id is required"}), 400
    
    ## make sure cosmos is configured
    cosmos_conversation_client = init_cosmosdb_client()
    if not cosmos_conversation_client:
        raise Exception("CosmosDB is not configured or not working")

    ## get the conversation object and the related messages from cosmos
    conversation = await cosmos_conversation_client.get_conversation(user_id, conversation_id)
    ## return the conversation id and the messages in the bot frontend format
    if not conversation:
        return jsonify({"error": f"Conversation {conversation_id} was not found. It either does not exist or the logged in user does not have access to it."}), 404
    
    # get the messages for the conversation from cosmos
    conversation_messages = await cosmos_conversation_client.get_messages(user_id, conversation_id)
        
    ## format the messages in the bot frontend format
    messages = [
        {
            'id': msg['id'],
            'role': msg['role'],
            'content': msg['content'],
            'createdAt': msg['createdAt'],
            'feedback': msg.get('feedback'),
            'hidden': msg.get('hidden')
        }
        for msg in conversation_messages
    ]

    await cosmos_conversation_client.cosmosdb_client.close()
    return jsonify({"conversation_id": conversation_id, "messages": messages}), 200

@bp.route("/history/rename", methods=["POST"])
async def rename_conversation():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)
    
    if not conversation_id:
        return jsonify({"error": "conversation_id is required"}), 400
    
    ## make sure cosmos is configured
    cosmos_conversation_client = init_cosmosdb_client()
    if not cosmos_conversation_client:
        raise Exception("CosmosDB is not configured or not working")
    
    ## get the conversation from cosmos
    conversation = await cosmos_conversation_client.get_conversation(user_id, conversation_id)
    if not conversation:
        return jsonify({"error": f"Conversation {conversation_id} was not found. It either does not exist or the logged in user does not have access to it."}), 404

    ## update the title
    title = request_json.get("title", None)
    if not title:
        return jsonify({"error": "title is required"}), 400
    conversation['title'] = title
    updated_conversation = await cosmos_conversation_client.upsert_conversation(conversation)

    await cosmos_conversation_client.cosmosdb_client.close()
    return jsonify(updated_conversation), 200

@bp.route("/history/delete_all", methods=["DELETE"])
async def delete_all_conversations():
    ## get the user id from the request headers
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']

    # get conversations for user
    try:
        ## make sure cosmos is configured
        cosmos_conversation_client = init_cosmosdb_client()
        if not cosmos_conversation_client:
            raise Exception("CosmosDB is not configured or not working")

        conversations = await cosmos_conversation_client.get_conversations(user_id, offset=0, limit=None)
        if not conversations:
            return jsonify({"error": f"No conversations for {user_id} were found"}), 404
        
        # delete each conversation
        for conversation in conversations:
            ## delete the conversation messages from cosmos first
            deleted_messages = await cosmos_conversation_client.delete_messages(conversation['id'], user_id)

            ## Now delete the conversation 
            deleted_conversation = await cosmos_conversation_client.delete_conversation(user_id, conversation['id'])
        await cosmos_conversation_client.cosmosdb_client.close()
        return jsonify({"message": f"Successfully deleted conversation and messages for user {user_id}"}), 200
    
    except Exception as e:
        logging.exception("Exception in /history/delete_all")
        return jsonify({"error": str(e)}), 500

@bp.route("/history/clear", methods=["POST"])
async def clear_messages():
    ## get the user id from the request headers
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    user_id = authenticated_user['user_principal_id']
    
    ## check request for conversation_id
    request_json = await request.get_json()
    conversation_id = request_json.get('conversation_id', None)

    try: 
        if not conversation_id:
            return jsonify({"error": "conversation_id is required"}), 400
        
        ## make sure cosmos is configured
        cosmos_conversation_client = init_cosmosdb_client()
        if not cosmos_conversation_client:
            raise Exception("CosmosDB is not configured or not working")

        ## delete the conversation messages from cosmos
        deleted_messages = await cosmos_conversation_client.delete_messages(conversation_id, user_id)

        return jsonify({"message": "Successfully deleted messages in conversation", "conversation_id": conversation_id}), 200
    except Exception as e:
        logging.exception("Exception in /history/clear_messages")
        return jsonify({"error": str(e)}), 500


@bp.route("/history/ensure", methods=["GET"])
async def ensure_cosmos():
    if not AZURE_COSMOSDB_ACCOUNT:
        return jsonify({"error": "CosmosDB is not configured"}), 404
    
    try:
        cosmos_conversation_client = init_cosmosdb_client()
        success, err = await cosmos_conversation_client.ensure()
        if not cosmos_conversation_client or not success:
            if err:
                return jsonify({"error": err}), 422
            return jsonify({"error": "CosmosDB is not configured or not working"}), 500
        
        await cosmos_conversation_client.cosmosdb_client.close()
        return jsonify({"message": "CosmosDB is configured and working"}), 200
    except Exception as e:
        logging.exception("Exception in /history/ensure")
        cosmos_exception = str(e)
        if "Invalid credentials" in cosmos_exception:
            return jsonify({"error": cosmos_exception}), 401
        elif "Invalid CosmosDB database name" in cosmos_exception:
            return jsonify({"error": f"{cosmos_exception} {AZURE_COSMOSDB_DATABASE} for account {AZURE_COSMOSDB_ACCOUNT}"}), 422
        elif "Invalid CosmosDB container name" in cosmos_exception:
            return jsonify({"error": f"{cosmos_exception}: {AZURE_COSMOSDB_CONVERSATIONS_CONTAINER}"}), 422
        else:
            return jsonify({"error": "CosmosDB is not working"}), 500


async def generate_title(conversation_messages):
    ## make sure the messages are sorted by _ts descending
    title_prompt = 'Summarize the conversation so far into a 4-word or less title. Do not use any quotation marks or punctuation. Respond with a json object in the format {{"title": string}}. Do not include any other commentary or description.'

    messages = [{'role': msg['role'], 'content': msg['content']} for msg in conversation_messages]
    messages.append({'role': 'user', 'content': title_prompt})

    try:
        azure_openai_client = init_openai_client(use_data=False)
        response = await azure_openai_client.chat.completions.create(
            model=AZURE_OPENAI_MODEL,
            messages=messages,
            temperature=1,
            max_tokens=64
        )
        
        title = json.loads(response.choices[0].message.content)['title']
        return title
    except Exception as e:
        return messages[-2]['content']
    

    
def chunkString(text, chunk_size,overlap):
    SENTENCE_ENDINGS = [".", "!", "?"]
    WORDS_BREAKS = list(reversed([",", ";", ":", " ", "(", ")", "[", "]", "{", "}", "\t", "\n"]))

    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)


    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        separators=SENTENCE_ENDINGS + WORDS_BREAKS,
        chunk_size=chunk_size, chunk_overlap=overlap)
    chunked_content_list = splitter.split_text(text)

    return chunked_content_list

async def collect_documents(filenames, container_name):
    try:
        container_client = init_container_client(container_name)
    except Exception as e:
        print(f"Error initializing container client: {e}")
        raise e
    docString = f'There are {len(filenames)} documents in this collection:'
    for index, filename in enumerate(filenames):
        blob_client = container_client.get_blob_client(filename)
        print(f"Downloading document: {filename}")
        try:
            download_stream = blob_client.download_blob()
            if filename.endswith('.docx'):
                print('reading docx')
                with io.BytesIO() as output_stream:
                    download_stream.readinto(output_stream)
                    doc = docx.Document(output_stream)
                    docString += f" START OF DOCUMENT {index+1} ({filename}):"
                    for para in doc.paragraphs:
                        docString += para.text
            elif filename.endswith('.pdf'):
                loader = AzStorageBlobReader(
                    account_url = f'https://{AZURE_STORAGE_ACCOUNT}.blob.core.windows.net',
                    container_name = container_name,
                    blob = filename,
                    credential = AZURE_STORAGE_KEY
                )
                documents = loader.load_data()
                documents = [doc.text for doc in documents if doc]
                docString += f" START OF DOCUMENT {index+1}, {filename}:"
                docString += ' '.join(documents)
            else:
                blob_bytes = download_stream.readall()
                docString += f" START OF DOCUMENT {index+1}, {filename}:"
                docString += blob_bytes.decode('utf-8')
        
        except Exception as e:
            print(f"Error downloading document: {e}")
    
    return docString



async def summarize_chunk(chunk: str, max_tokens: int, prompt = 'Produce a detailed summary of the following including all key concepts and takeaways, if it is a guide or a help piece make sure you include a summary of the main actionable steps:') -> str:
    azure_openai_client = init_openai_client(use_data=False)
    response = await azure_openai_client.chat.completions.create(
        model=AZURE_OPENAI_MODEL,
        messages=[{'role': 'system', 'content': AZURE_OPENAI_SYSTEM_MESSAGE},
                  {'role': 'user', 'content': f'{prompt} {chunk}'}],
        temperature=1,
        max_tokens=max_tokens
    )
    return response.choices[0].message.content


@bp.websocket('/documentsummary/refine')
async def refineProgress():
    try:
        
        abort_flag = False

        async def listen_for_abort():
            nonlocal abort_flag
            while not abort_flag:
                data = await websocket.receive()
                data = json.loads(data)
                if data.get('command') == 'abort':
                    abort_flag = True

        asyncio.create_task(listen_for_abort())

        while True:
            data = await websocket.receive()
            data = json.loads(data)
            filenames = data['filenames']
            prompt = data['prompt']
            container_name = data['container']
            if not prompt == '':
                prompt = 'In your answer adhere to the following instruction: ' + prompt + '. Here is the document: '
            try:
                document = await collect_documents(filenames, container_name)
            except Exception as e:
                print(f"Error collecting documents: {e}")
                return
            try:
                chunks = chunkString(document, 3500, 100)
                combined_summaries = ""
                for chunk in chunks:
                    if abort_flag:
                        return
                    index = chunks.index(chunk)
                    if index == len(chunks) - 1:
                        break
                    await websocket.send(f"{chunks.index(chunk)+1}/{len(chunks)}")
                    combined_summaries = await summarize_chunk(combined_summaries + chunk, 4000)
                final_prompt = f'Tell me in as much detail as possible what the following document is about. Make sure your answer includes the concepts, themes, priciples and methods that are covered. {prompt}'
                await websocket.send(f"done:{final_prompt} {combined_summaries} {chunks[len(chunks)-1]}")
            except Exception as e:
                print(f"Error  reading document: {e}")
    finally:
        await websocket.close()
    
async def document_summary_reduce_internal(document, prompt):
    if not prompt == '':
        prompt = 'In your answer adhere to the following instruction: ' + prompt + '. Here is the document: '
    try:
        chunks = chunkString(document, 4000, 100)
        chunk_summaries = await asyncio.gather(*[summarize_chunk(chunk, round(4000/len(chunks))) for chunk in chunks])
        combined_summaries = " ".join(chunk_summaries)
        final_prompt = f'I have a document that I have broken into chunks, the folliwng is the summary of each chunk. Tell me in as much detail as possible what the document is about. Make sure your answer includes the concepts, themes, priciples and methods that are covered. {prompt}'
        return jsonify({"response": f'{final_prompt} {combined_summaries}'}), 200
    except Exception as e:
        print(f"Error reducing document: {e}")
        return jsonify({"error": f"Error reading document: {e}"}), 500
    

@bp.route("/documentsummary/reduce", methods=["POST"])
async def documentsummary():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    storage_container_name = authenticated_user['user_principal_id']
    data = await request.json
    filenames = data.get("filenames", [])
    prompt = data.get("prompt", "")

    if len(filenames) == 0:
        return jsonify({"error": "Filenames not found"}), 400
    docString = await collect_documents(filenames, storage_container_name)
    return await document_summary_reduce_internal(docString, prompt)



@bp.route("/create_search_index", methods=["POST"])
async def create_search_index():
    print("Python HTTP trigger function processed a request for CreateSearchIndex.")
    # 'Python HTTP trigger function processed a request for CreateSearchIndex.'

    # """
    #     If the index already exists, then this function will update it and won't create a new one from scratch.
    #     If the updates are not compatible with the existing index, it will throw an error.

    #     HTTP call body example:
    #         {
    #             "search_index_name": "test-ingrid",
    #             "synonym_map_name": "test-synonyms"
    #         }
    # """

    try:
        # get the more static environment variables from the function app
        search_service_name = AZURE_SEARCH_SERVICE
        search_service_key = AZURE_SEARCH_KEY
        search_index_name = AZURE_SEARCH_INDEX

        # get Azure credentials
        # credential_MI = DefaultAzureCredential() # use this when MI is set up for Azure AI Search
        credential = AzureKeyCredential(search_service_key) # use this for free tier Azure AI Search (MI not supported)
        # connect to Azure Cognitive Search resource
        service_endpoint = f"https://{search_service_name}.search.windows.net"
        index_client = SearchIndexClient(service_endpoint, credential)

        # define the fields we want the index to contain
        fields = [
            SimpleField(name="id", type=SearchFieldDataType.String, key=True),
            SearchableField(name="user_id", type=SearchFieldDataType.String, filterable=True),
            SimpleField(name="doc_url", type=SearchFieldDataType.String),
            SearchableField(name="filename", type=SearchFieldDataType.String, filterable=True),
            SearchableField(name="content", type=SearchFieldDataType.String),
            SearchField(
                name="content_vector", 
                type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                searchable=True, 
                vector_search_dimensions=1536, 
                vector_search_profile_name="ll-vector-config"
            ),
            SearchableField(name="llamaindex_metadata", type=SearchFieldDataType.String),
            SearchableField(name="llamaindex_doc_id", type=SearchFieldDataType.String, filterable=True)
        ]

        # For vector search, we want to use the HNSW (Hierarchical Navigable Small World)
        # algorithm (a type of approximate nearest neighbor search algorithm) with cosine distance.
        vector_search = VectorSearch(
            profiles=[
                VectorSearchProfile(
                    name="ll-vector-config", 
                    algorithm_configuration_name="ll-algorithms-config"
                )
            ],
            algorithms=[
                HnswAlgorithmConfiguration(
                    name="ll-algorithms-config",
                    kind="hnsw",
                    parameters=HnswParameters(metric="cosine")
                )
            ]
        )

        # add a scoring profile to treat the content and vector with higher importance than question
        scoring_profiles = [
            {
                "name": "Hybrid",
                "functionAggregation": "sum",
                "text": {
                    "weights": {
                    "content": 100,
                    "content_vector": 100
                    }
                },
                "functions": []
            }
        ]

        # Put the search index together
        index = SearchIndex(
            name=search_index_name,
            fields=fields,
            vector_search=vector_search,
            scoring_profiles=scoring_profiles
        )

        # Create the index
        index_client.create_or_update_index(index)
        print(f"Index named \"{search_index_name}\" has been created successfully.")
        return jsonify(
            {"success": f"Index named \"{search_index_name}\" has been created successfully."}
            ,200
        )
    
    except Exception as ex:
        print(f"Processing failed. Exception: {ex}")
        return jsonify(
             {"error": f"Processing failed. Exception: {ex}"}, 500
        )


@bp.route("/upload_documents", methods=["POST"])
async def upload_documents():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    storage_container_name = authenticated_user['user_principal_id']
    uploadedFiles = []
    try:

        files = await request.files
        if 'file' not in files:
            print("No files part in the request")
            return jsonify({"error": "No files part in the request"}), 400
        file_storage_list = files.getlist('file')

        if not file_storage_list:
            print("No files selected for uploading")
            return jsonify({"error": "No files selected for uploading"}), 400
        
        for doc in file_storage_list:

            file_name = doc.filename
            storage_account_url = f"https://{AZURE_STORAGE_ACCOUNT}.blob.core.windows.net/"

            blob_service_client = BlobServiceClient(
                account_url=storage_account_url, 
                credential=AZURE_STORAGE_KEY
            )

            container_client = blob_service_client.get_container_client(storage_container_name)
            try:
                container_client.create_container()
                print(f"Container '{storage_container_name}' created.")
            except Exception as e:
                if "ContainerAlreadyExists" in str(e):
                    print(f"Container '{storage_container_name}' already exists.")
                else:
                    raise e
            
            blob_client = blob_service_client.get_blob_client(
                container=storage_container_name,
                blob=file_name
            )
            blob_client.upload_blob(doc.stream, overwrite=True)
            uploadedFiles.append((file_name, blob_client.url))
            print(f"Successfully uploaded {file_name} at location {blob_client.url}.")
        result = await addDocuments(uploadedFiles, storage_container_name)
        print(f"Uploading Documents successful. Result: {result}")
        return jsonify({"success": f"Uploading Documents successful. Result: {result}"}, 200)
    except Exception as ex:
        print(f"Uploading Documents failed. Exception: {ex}")
        return jsonify({"error": f"Uploading Documents failed. Exception: {ex}"}, 500)


async def addDocuments(document_tuples, container_name: str):
    print('Python HTTP trigger function processed a request for AddDocuments.')

    def convert_doc(documents, blob_name, url):
        docs = []
        full_text = []
        doc_url_list = ['doc_url', 'user_id']
        # for every page in the document
        for document in documents:
            
            document.text = document.text.replace("\n",".  ").replace('..', '.')
            document.text = re.sub(r"\s+", " ", document.text)
            # extract and merge the text from individual pages
            full_text.append(document.text)

        # create a new Document with metadata fields and all text in one Document
        docs.append(Document(

            # make sure ID of the document is the pdf name without extension
            id_ = re.sub(r'[^\w&=.-]', '_', blob_name).strip('_').replace(' ', '_').replace('.pdf','').replace('.PDF',''),

            # join the text of the PDF into a single string
            text = ' '.join(full_text),

            # add metadata fields
            metadata = {
                #'title': blob_name.replace('.pdf','').replace('.PDF',''),
                'user_id': container_name,
                'doc_url': url,
                'filename': blob_name,
            },
            # don't use the fields added as metadtata in the embedding/llm processes
            excluded_embed_metadata_keys = doc_url_list,
            excluded_llm_metadata_keys = doc_url_list
        ))
        return docs
    
    try:
        storage_account_name = AZURE_STORAGE_ACCOUNT
        storage_account_key = AZURE_STORAGE_KEY
        embedding_deployment_name = AZURE_OPENAI_EMBEDDING_NAME
        embedding_model_name = AZURE_OPENAI_EMBEDDING_MODEL_NAME
        llm_deployment_name = AZURE_OPENAI_MODEL
        llm_model_name = AZURE_OPENAI_MODEL_NAME
        api_key = AZURE_OPENAI_KEY
        api_version = AZURE_OPENAI_PREVIEW_API_VERSION
        azure_endpoint = AZURE_OPENAI_ENDPOINT
        chunk_size = 1000
        chunk_overlap = 100


        # get Azure credentials

        # connect to an existing Azure Cognitive Search index
        search_client = init_search_client()


        container_client = init_container_client(container_name)



        # define the llm model (using Azure OpenAI gpt-35-turbo model)
        llm = AzureOpenAI(
            model=llm_model_name,
            deployment_name=llm_deployment_name,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
        )


        # define the embedding model (using Azure OpenAI text-embedding-ada-002 model)
        embed_model = AzureOpenAIEmbedding(
            model=embedding_model_name,
            deployment_name=embedding_deployment_name,
            api_key=api_key,
            azure_endpoint=azure_endpoint,
            api_version=api_version,
        )


        Settings.llm = llm
        Settings.embed_model = embed_model


        final_nodes = []

        # process policy documents in each location
        for doc in document_tuples:
            blob_name = doc[0]
            url = doc[1]
            
            if blob_name.endswith('.docx'):
                blob_client = container_client.get_blob_client(blob_name)
                print(f"Downloading document: {blob_name}")
                try:
                    download_stream = blob_client.download_blob()
                    with io.BytesIO() as output_stream:
                        download_stream.readinto(output_stream)
                        doc = docx.Document(output_stream)
                        documents = []
                        for para in doc.paragraphs:
                            documents.append(para)
                except Exception as e:
                    print(f"Error downloading document: {e}")

            else:
                print(f'Processing {blob_name} with url {url}.')
                # create the loader

                loader = AzStorageBlobReader(
                    account_url = f'https://{storage_account_name}.blob.core.windows.net',
                    container_name = container_name,
                    blob = blob_name,
                    credential = storage_account_key # replace this with DefaultAzureCredential() once MI is set up
                )

                # load in document
                documents = loader.load_data()

                # convert documents to the correct format for us
            docs = convert_doc(documents, blob_name, url)

            # define a node parser (specify chunk size and overalap)
            node_parser = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

            # split the documents into nodes (aka chunks)
            nodes = node_parser.get_nodes_from_documents(docs, show_progress=False)

            # change the node ID to document name + chunk/node number
            chunk_num = 1
            for node in nodes:
                node.id_ = re.sub(r'[^a-zA-Z0-9_]', '_', blob_name).strip('_').replace(' ', '_').replace('_pdf','').replace('_PDF','').replace('___', '_') + str(chunk_num)
                chunk_num = chunk_num + 1

            # generate embeddings for each node
            for node in nodes:
                node_embedding = embed_model.get_text_embedding(
                    node.get_content(metadata_mode="all")
                )
                node.embedding = node_embedding

            # add all the nodes to a final list that will be added to the search index
            final_nodes.append(nodes)

        # flatten the list of nodes
        flat_nodes_list = [item for node in final_nodes for item in node]

        # define additional fields to be added to the search index 
        # (needs to match the fields added in CreateSearchIndex)
        metadata_fields = {
            'user_id': 'user_id',
            'doc_url': 'doc_url',
            'filename': 'filename'
        }


        # define the vector store we will be connecting to 
        # (needs to match the fields added in CreateSearchIndex)
        vector_store = AzureAISearchVectorStore(
            search_or_index_client=search_client,
            filterable_metadata_field_keys=metadata_fields,
            index_management=IndexManagement.VALIDATE_INDEX,
            id_field_key="id",
            chunk_field_key="content",
            embedding_field_key="content_vector",
            embedding_dimensionality=1536,
            metadata_string_field_key="llamaindex_metadata",
            doc_id_field_key="llamaindex_doc_id"
        )

        # set the storage context
        storage_context = StorageContext.from_defaults(vector_store=vector_store)

        
        # load the index
        index = VectorStoreIndex.from_documents(
            [],
            storage_context=storage_context,
        )


        # add the nodes to the vector store/search index
        index.insert_nodes(flat_nodes_list)
        ##ERROR HAPPENS HERE : Processing failed. Exception: Object of type set is not JSON serializable




        # get the ID-s of the documents added to the search index
        node_ids = ', '.join([node.id_ for node in flat_nodes_list])

        print(f"Documents added to the search index. Document IDs: {node_ids}")
        return f"Documents added to the search index. Document IDs: {node_ids}"
    
    except Exception as ex:

        return f"Document addition failed. Exception: {ex}"


@bp.route("/get_documents", methods=["GET"])
async def get_documents():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    container_client = init_container_client(authenticated_user['user_principal_id'])
    try:
        blob_list = container_client.list_blob_names()
        blob_names = [blob for blob in blob_list]
        return jsonify(blob_names), 200
    except Exception as ex:
        return jsonify({"error": f"Failed to get documents. Exception: {ex}"}), 500

@bp.route("/delete_documents", methods=["POST"])
async def delete_documents():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    container_client = init_container_client(authenticated_user['user_principal_id'])
    request_json = await request.get_json()
    requested_blobs = request_json.get('filenames', [])
    try:
        search_client = init_search_client()
        blob_list = container_client.list_blob_names()
        for blob in requested_blobs:
            print(f"Deleting {blob}")

            try:
                container_client.delete_blob(blob)
            except Exception as ex:
                print(f"Failed to delete {blob} from storage. Exception: {ex}")
            try:
                results = search_client.search(search_text=blob, search_fields=["filename"])
                for doc in results:
                    print(doc['id'])
                    if doc['filename'] == blob:
                        search_client.delete_documents(documents=[{"id": doc['id']}])
                        print(f"Deleted {doc['id']} from search index")
                    else:
                       Exception('Document not found in search index')
            except Exception as ex:
                print(f"Failed to delete {blob} from search index. Exception: {ex}")
        return jsonify({"success": "All documents deleted"}), 200
    except Exception as ex:
        return jsonify({"error": f"Failed to delete documents. Exception: {ex}"}), 500

@bp.route("/get_user_id", methods=["GET"])
async def get_user_id():
    authenticated_user = get_authenticated_user_details(request_headers=request.headers)
    return jsonify(authenticated_user['user_principal_id']), 200

app = create_app()